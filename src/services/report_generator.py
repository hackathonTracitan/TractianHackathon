import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

import pandas as pd

def generate_report_file(messages: list, df: pd.DataFrame) -> bytes:
    """Converts a chat to a Word document and returns the document as a byte string."""

    document = Document()

    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"

    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    image_path = Path(__file__).parent.parent.parent / "images" / "tractian.png"
    run.add_picture(str(image_path), width=Inches(1.5))
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_heading("Relatório de especificação de máquina", 0)
    document.add_heading(
        f"Data e hora: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        level=1,
    )
    document.add_heading("Especificações:\n", level=1)
    p = document.add_paragraph()

    for message in messages:
        p.add_run(message)

    t = document.add_table(df.shape[0]+1, df.shape[1])
    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    word_file_io = BytesIO()
    document.save(word_file_io)
    word_file_io.seek(0)
    return word_file_io
